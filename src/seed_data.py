import re, glob
import streamlit as st
from pathlib import Path
from docx import Document
from langchain.schema import Document as LC_Document
from langchain.retrievers import PineconeHybridSearchRetriever
from pinecone_text.sparse import BM25Encoder
from langchain_pinecone import PineconeVectorStore
from pinecone import Pinecone, ServerlessSpec
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.document_loaders import DirectoryLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter

PINECONE_API_KEY = st.secrets["PINECONE_API_KEY"]
if not PINECONE_API_KEY:
    raise ValueError("PINECONE_API_KEY not found in environment variables")

def read_docx(file_path: str) -> list[str]:
    doc = Document(file_path)
    content = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraph = paragraph.text.lower()
            content.append(paragraph.strip())
    return content

def extract_sections(text: list[str]) -> tuple:
    title = f"{text[0].capitalize()} {text[1].capitalize()}"
    sections = {}
    current_section = None
    for line in text:
        match = re.match(r"^\s*(điều \d+\. .+)$", line, re.IGNORECASE)
        if match:
            current_section = match.group(1)
            sections[current_section] = []
        elif current_section:
            sections[current_section].append(line)
    return title, sections

def article_chunk(path: list[str]) -> list[LC_Document]:
    file_list = glob.glob(path)
    basename_list = [Path(p).stem for p in file_list]
    documents = []
    for i, path in enumerate(file_list):
        loader = read_docx(path)
        title, sections = extract_sections(loader)
        keys = list(sections.keys())
        values = list(sections.values())
        for j in range(len(sections)):
            page_content = values[j]
            if isinstance(page_content, list):
                page_content = " ".join(page_content)
            document = LC_Document(
                page_content=page_content,
                metadata={"source": basename_list[i], 
                          "title": title,
                          "article": keys[j].split(".")[0].strip(),
                          "article_title": keys[j].split(".")[1].strip()}
            )
            documents.append(document)
    return documents

def call_index(index_name: str, 
               dimention: int, 
               metric: str):
    pc = Pinecone(api_key=PINECONE_API_KEY)
    if index_name not in pc.list_indexes().names():
        pc.create_index(name=index_name,
                        dimension=dimention,
                        metric=metric,
                        spec=ServerlessSpec(cloud='aws', region='us-east-1'))
    return pc.Index(index_name)

def pinecone_hybrid_retriever(index_name: str,
                  path: str,  
                  dimentions: int = 768, 
                  metric: str = "dotproduct",
                  embeddings_name: str = "hiieu/halong_embedding",
                  article_chunk: bool = True) -> PineconeHybridSearchRetriever:
    embeddings = HuggingFaceEmbeddings(model_name=embeddings_name)
    if article_chunk == False:    
        loader = DirectoryLoader(path=path, glob="**/*.docx")
        documents = loader.load()
    else:
        documents = article_chunk(path)
    index = call_index(index_name, dimentions, metric)
    bm25_encoder = BM25Encoder().default()
    
    splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100)
    docs = splitter.split_documents(documents)
    list_docs = [doc.page_content for doc in docs]
    list_metadatas = [doc.metadata for doc in docs]
    retriever = PineconeHybridSearchRetriever(embeddings=embeddings, 
                                              sparse_encoder=bm25_encoder, 
                                              index=index,
                                              alpha=0.7)
    retriever.add_text(texts= list_docs, metadatas= list_metadatas)
    return retriever

@st.cache_resource
def get_retriever(index_name: str, 
                    embeddings_name: str = "hiieu/halong_embedding") -> PineconeVectorStore:
    index = call_index(index_name, 768, "dotproduct")
    embeddings = HuggingFaceEmbeddings(model_name = embeddings_name)
    bm25_encoder = BM25Encoder().default()
    retriever = PineconeHybridSearchRetriever(index=index, embeddings=embeddings, sparse_encoder=bm25_encoder)
    # vectorstore = PineconeVectorStore(index=index, embedding=embeddings)
    return retriever

def main():
    pinecone_hybrid_retriever('hybrid-rag', 'data', 768, 'dotproduct')
    
if __name__ == "__main__":
    main()